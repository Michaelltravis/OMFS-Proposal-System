"""
Document Export Service for generating Word documents from proposal content
"""
from typing import Optional, List
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from bs4 import BeautifulSoup
from io import BytesIO
import re


class DocumentExportService:
    """Service for exporting proposal sections and content to Word documents"""

    def __init__(self):
        self.default_font = "Calibri"
        self.default_size = Pt(11)

    def export_section_to_docx(
        self,
        section_title: str,
        section_contents: List[dict],
        formatting_instructions: Optional[str] = None
    ) -> BytesIO:
        """
        Export a proposal section to a Word document

        Args:
            section_title: Title of the section
            section_contents: List of content items with 'title' and 'content' (HTML)
            formatting_instructions: Optional Claude-generated formatting instructions

        Returns:
            BytesIO object containing the Word document
        """
        doc = Document()

        # Set default styles
        self._set_default_styles(doc)

        # Add section title as main heading
        heading = doc.add_heading(section_title, level=1)
        heading.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

        # Add each content item
        for content_item in section_contents:
            if content_item.get('title'):
                # Add content item title as subheading
                doc.add_heading(content_item['title'], level=2)

            # Convert HTML content to Word paragraphs
            self._html_to_docx(doc, content_item.get('content', ''))

            # Add spacing between content items
            doc.add_paragraph()

        # Apply custom formatting if provided
        if formatting_instructions:
            self._apply_formatting_instructions(doc, formatting_instructions)

        # Save to BytesIO
        buffer = BytesIO()
        doc.save(buffer)
        buffer.seek(0)

        return buffer

    def _set_default_styles(self, doc: Document):
        """Set default document styles"""
        styles = doc.styles

        # Style for normal paragraphs
        if 'Normal' in styles:
            normal_style = styles['Normal']
            normal_style.font.name = self.default_font
            normal_style.font.size = self.default_size
            normal_style.paragraph_format.space_after = Pt(6)
            normal_style.paragraph_format.line_spacing = 1.15

    def _html_to_docx(self, doc: Document, html_content: str):
        """
        Convert HTML content to Word document paragraphs

        Args:
            doc: Document object to add content to
            html_content: HTML string to convert
        """
        if not html_content:
            return

        soup = BeautifulSoup(html_content, 'html.parser')

        for element in soup.find_all(recursive=False):
            self._process_element(doc, element)

    def _process_element(self, doc: Document, element, parent_paragraph=None):
        """Process a single HTML element"""

        if element.name in ['h1', 'h2', 'h3', 'h4', 'h5', 'h6']:
            # Convert headings
            level = int(element.name[1])
            doc.add_heading(element.get_text(), level=level)

        elif element.name == 'p':
            # Process paragraph
            para = doc.add_paragraph()
            self._process_text_content(para, element)

        elif element.name in ['ul', 'ol']:
            # Process lists
            for li in element.find_all('li', recursive=False):
                para = doc.add_paragraph(style='List Bullet' if element.name == 'ul' else 'List Number')
                self._process_text_content(para, li)

        elif element.name == 'table':
            # Process table
            self._process_table(doc, element)

        elif element.name == 'br':
            # Line break
            if parent_paragraph:
                parent_paragraph.add_run().add_break()

        else:
            # For other elements, extract text
            text = element.get_text()
            if text.strip():
                para = doc.add_paragraph(text)

    def _process_text_content(self, paragraph, element):
        """Process text content with formatting (bold, italic, etc.)"""

        for child in element.children:
            if isinstance(child, str):
                # Plain text
                text = child.strip()
                if text:
                    paragraph.add_run(text)
            else:
                # Handle formatting tags
                text = child.get_text()
                run = paragraph.add_run(text)

                if child.name in ['strong', 'b']:
                    run.bold = True
                elif child.name in ['em', 'i']:
                    run.italic = True
                elif child.name == 'u':
                    run.underline = True

    def _process_table(self, doc: Document, table_element):
        """Process HTML table and convert to Word table"""

        # Count rows and columns
        rows = table_element.find_all('tr')
        if not rows:
            return

        max_cols = max(len(row.find_all(['td', 'th'])) for row in rows)

        # Create Word table
        word_table = doc.add_table(rows=len(rows), cols=max_cols)
        word_table.style = 'Light Grid Accent 1'

        # Fill table
        for row_idx, row in enumerate(rows):
            cells = row.find_all(['td', 'th'])
            for col_idx, cell in enumerate(cells):
                word_table.rows[row_idx].cells[col_idx].text = cell.get_text().strip()

    def _apply_formatting_instructions(self, doc: Document, instructions: str):
        """
        Apply formatting instructions from Claude to the document

        Args:
            doc: Document object
            instructions: String containing formatting instructions
        """
        # This is a placeholder for Claude-driven formatting
        # In a real implementation, Claude could provide structured formatting
        # instructions that are parsed and applied here

        # Example: Claude could return JSON with specific formatting rules
        # For now, we'll keep the default formatting
        pass

    def export_full_proposal_to_docx(
        self,
        proposal_title: str,
        sections: List[dict],
        formatting_instructions: Optional[str] = None
    ) -> BytesIO:
        """
        Export a full proposal with all sections to Word

        Args:
            proposal_title: Title of the proposal
            sections: List of section dictionaries with title and contents
            formatting_instructions: Optional Claude-generated formatting instructions

        Returns:
            BytesIO object containing the Word document
        """
        doc = Document()
        self._set_default_styles(doc)

        # Add title page
        title = doc.add_heading(proposal_title, level=0)
        title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        doc.add_page_break()

        # Add each section
        for section in sections:
            # Section title
            doc.add_heading(section.get('title', 'Untitled Section'), level=1)

            # Section contents
            for content_item in section.get('contents', []):
                if content_item.get('title'):
                    doc.add_heading(content_item['title'], level=2)

                self._html_to_docx(doc, content_item.get('content', ''))
                doc.add_paragraph()

            # Page break after each section (except last)
            if section != sections[-1]:
                doc.add_page_break()

        # Apply custom formatting if provided
        if formatting_instructions:
            self._apply_formatting_instructions(doc, formatting_instructions)

        # Save to BytesIO
        buffer = BytesIO()
        doc.save(buffer)
        buffer.seek(0)

        return buffer


# Create singleton instance
document_export_service = DocumentExportService()
